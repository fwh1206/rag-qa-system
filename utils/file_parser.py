"""文件解析工具：把 PDF/Word/文本/Markdown/Excel 提取为纯文本，并提供中文感知切分。"""

import json

from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader


def split_text(text: str, chunk_size: int = 400, chunk_overlap: int = 50) -> list[str]:
    """基于 RecursiveCharacterTextSplitter 的中文感知分片。

    分隔符按中文阅读习惯排列优先级：段落 > 换行 > 句号/叹号/问号/分号/逗号 > 空格。
    参数非法时回退到默认值，保证调用方拿到的切片长度可控。
    """
    if not text or not text.strip():
        return []
    if chunk_size <= 0:
        chunk_size = 400
    if chunk_overlap < 0 or chunk_overlap >= chunk_size:
        chunk_overlap = max(0, chunk_size // 8)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
        keep_separator=True,
        length_function=len,
    )
    return [chunk.strip() for chunk in splitter.split_text(text) if chunk and chunk.strip()]


def _read_txt(file_path: str) -> str:
    """读取文本文件，优先 UTF-8（含 BOM），失败回退 GB18030，再失败用替换模式兜底。"""
    with open(file_path, "rb") as f:
        raw = f.read()
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _read_docx(file_path: str) -> str:
    """提取 docx 段落与表格文本，表格按行用竖线拼接，便于后续检索。"""
    doc = Document(file_path)
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _flatten_json(value, prefix: str, lines: list[str]):
    """把 JSON 递归拍平成「路径: 值」文本，便于向量检索。"""
    if isinstance(value, dict):
        for key, item in value.items():
            _flatten_json(item, _join_json_path(prefix, key), lines)
    elif isinstance(value, list):
        for index, item in enumerate(value):
            _flatten_json(item, _join_json_path(prefix, f"[{index}]"), lines)
    else:
        lines.append(f"{prefix}: {value}")


def _join_json_path(prefix: str, part: str) -> str:
    if not prefix:
        return part
    if part.startswith("["):
        return prefix + part
    return prefix + "." + part


def _read_json(file_path: str) -> str:
    data = json.loads(_read_txt(file_path))
    lines: list[str] = []
    _flatten_json(data, "", lines)
    return "\n".join(lines)


def _read_doc(file_path: str) -> str:
    """解析旧版 .doc：优先使用本机 Word，其次 WPS，均不可用时给出明确提示。"""
    try:
        from win32com.client import Dispatch
    except ImportError as exc:
        raise ValueError("解析 .doc 需要本机安装 pywin32，请先 pip install pywin32") from exc
    errors = []
    for prog_id in ("Word.Application", "KWPS.Application", "kwps.application"):
        try:
            app = Dispatch(prog_id)
            app.Visible = False
            try:
                doc = app.Documents.Open(file_path, ReadOnly=True)
                text = doc.Content.Text
                doc.Close(False)
                return text
            finally:
                app.Quit()
        except Exception as exc:
            errors.append(f"{prog_id}: {exc}")
    raise ValueError("解析 .doc 失败，请安装 Word/WPS 或转换为 .docx 后重新上传；" + "；".join(errors))


def _read_excel(file_path: str, suffix: str) -> str:
    """读取 Excel：xlsx 用 openpyxl，xls 用 xlrd，输出为「工作表标题 + 行文本」。"""
    if suffix == ".xlsx":
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True, data_only=True)
        lines = []
        try:
            for sheet in wb.worksheets:
                lines.append(f"# 工作表：{sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                    if cells:
                        lines.append(" | ".join(cells))
        finally:
            wb.close()
        return "\n".join(lines)

    import xlrd

    book = xlrd.open_workbook(file_path)
    lines = []
    for sheet in book.sheets():
        lines.append(f"# 工作表：{sheet.name}")
        for r in range(sheet.nrows):
            cells = []
            for c in range(sheet.ncols):
                value = sheet.cell_value(r, c)
                if value not in ("", None):
                    cells.append(str(value))
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def read_file(file_path: str, suffix: str) -> str:
    """按后缀分发到对应解析器，统一返回纯文本。"""
    suffix = (suffix or "").lower()
    if suffix == ".pdf":
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        full_text = "\n".join(pages)
    elif suffix in (".txt", ".md"):
        full_text = _read_txt(file_path)
    elif suffix == ".json":
        full_text = _read_json(file_path)
    elif suffix == ".docx":
        full_text = _read_docx(file_path)
    elif suffix == ".doc":
        full_text = _read_doc(file_path)
    elif suffix in (".xlsx", ".xls"):
        full_text = _read_excel(file_path, suffix)
    else:
        raise ValueError(f"不支持的文件类型：{suffix}")
    return full_text.strip()
